"""
=======================================================
  GERADOR DE APR — API FastAPI (Railway)
  Endpoints: POST /processar  |  GET /stream/{job_id}
             GET /download/{job_id}/{tipo}
=======================================================
  pip install fastapi uvicorn python-multipart openai
              reportlab openpyxl pypdf python-docx
              pandas Pillow aiofiles
=======================================================
"""

import os
import sys
import json
import uuid
import base64
import re as _re
import asyncio
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, Any

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, FileResponse, Response

# ─────────────────────────────────────────────────────
#  CONFIG
# ─────────────────────────────────────────────────────
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")

TEMP_DIR = Path(tempfile.gettempdir()) / "apr_jobs"
TEMP_DIR.mkdir(exist_ok=True)

jobs: Dict[str, Dict[str, Any]] = {}

# ─────────────────────────────────────────────────────
#  APP
# ─────────────────────────────────────────────────────
app = FastAPI(title="APR Generator API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition", "Content-Type"],
)


# ══════════════════════════════════════════════════════
#  LEITORES DE ARQUIVO
# ══════════════════════════════════════════════════════

def ler_pdf(caminho):
    from pypdf import PdfReader
    return "".join(p.extract_text() or "" for p in PdfReader(caminho).pages)

def ler_excel(caminho):
    import pandas as pd
    ext = Path(caminho).suffix.lower()
    sheets = pd.read_excel(caminho, sheet_name=None, engine="xlrd" if ext == ".xls" else None)
    texto = ""
    for nome, df in sheets.items():
        texto += f"\n=== ABA: {nome} ===\n" + df.fillna("").to_string(index=False) + "\n"
    return texto

def ler_word(caminho):
    from docx import Document
    doc = Document(caminho)
    linhas = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                linhas.append(" | ".join(cells))
    return "\n".join(linhas)

def ler_csv(caminho):
    import pandas as pd
    return pd.read_csv(caminho, sep=None, engine="python").fillna("").to_string(index=False)

def ler_txt(caminho):
    return open(caminho, encoding="utf-8", errors="ignore").read()

def ler_json(caminho):
    return json.dumps(json.load(open(caminho, encoding="utf-8")), ensure_ascii=False, indent=2)

def ler_imagem(caminho):
    ext_mime = {".jpg":"image/jpeg",".jpeg":"image/jpeg",".png":"image/png",
                ".gif":"image/gif",".webp":"image/webp",".bmp":"image/bmp"}
    mime = ext_mime.get(Path(caminho).suffix.lower(), "image/jpeg")
    b64  = base64.b64encode(open(caminho,"rb").read()).decode()
    texto = ""
    try:
        import pytesseract
        from PIL import Image
        texto = pytesseract.image_to_string(Image.open(caminho), lang="por+eng")
    except Exception:
        pass
    return texto, b64, mime

def carregar_arquivo(caminho):
    ext = Path(caminho).suffix.lower()
    tipo_map = {
        ".pdf":"pdf",".xlsx":"excel",".xls":"excel",".xlsm":"excel",
        ".docx":"word",".doc":"word",".csv":"csv",".tsv":"csv",
        ".txt":"txt",".md":"txt",".log":"txt",".json":"json",
        ".jpg":"imagem",".jpeg":"imagem",".png":"imagem",
        ".gif":"imagem",".webp":"imagem",".bmp":"imagem",
    }
    tipo = tipo_map.get(ext, "txt")
    resultado = {"tipo": tipo, "texto": "", "base64": "", "mime": ""}
    if   tipo == "pdf":    resultado["texto"] = ler_pdf(caminho)
    elif tipo == "excel":  resultado["texto"] = ler_excel(caminho)
    elif tipo == "word":   resultado["texto"] = ler_word(caminho)
    elif tipo == "csv":    resultado["texto"] = ler_csv(caminho)
    elif tipo == "json":   resultado["texto"] = ler_json(caminho)
    elif tipo == "imagem":
        t, b, m = ler_imagem(caminho)
        resultado["texto"], resultado["base64"], resultado["mime"] = t, b, m
    else:
        resultado["texto"] = ler_txt(caminho)
    return resultado


# ══════════════════════════════════════════════════════
#  CHUNKING INTELIGENTE
# ══════════════════════════════════════════════════════

MAX_CHARS_EXTRACAO = 80000
MAX_CHARS_GERACAO  = 60000

def _preparar_texto(texto: str, limite: int) -> str:
    if len(texto) <= limite:
        return texto
    parte1 = int(limite * 0.60)
    parte2 = limite - parte1
    aviso = f"\n\n[... DOCUMENTO TRUNCADO — {len(texto)} chars totais ...]\n\n"
    return texto[:parte1] + aviso + texto[-parte2:]


# ══════════════════════════════════════════════════════
#  EXTRAÇÃO DE INFO
# ══════════════════════════════════════════════════════

def extrair_info(conteudo: dict, modelo: str) -> dict:
    from openai import OpenAI

    prompt_sys = """Você é especialista em segurança do trabalho brasileiro.
Analise o documento e extraia os dados para uma APR.

Retorne SOMENTE JSON válido, sem markdown:
{
  "empresa": "...",
  "responsavel": "...",
  "num_apr": "...",
  "atividade": "descrição da atividade principal",
  "data_inicio": "DD/MM/AAAA ou vazio",
  "data_termino": "DD/MM/AAAA ou vazio",
  "etapas": [
    {
      "numero": 1,
      "etapa": "Nome da etapa",
      "ferramentas": "Ferramentas/equipamentos",
      "perigos": "Perigos já mapeados ou vazio",
      "como_prevenir": "Medidas ou vazio"
    }
  ]
}
REGRAS:
- Extraia SOMENTE dados explícitos para campos administrativos
- Identifique TODAS as etapas/fases do documento
- Se o doc não tiver riscos mapeados, deixe perigos e como_prevenir vazios"""

    texto_completo = _preparar_texto(conteudo["texto"], MAX_CHARS_EXTRACAO)

    if conteudo["tipo"] == "imagem" and conteudo["base64"]:
        msgs = [{"role":"user","content":[
            {"type":"image_url","image_url":{"url":f"data:{conteudo['mime']};base64,{conteudo['base64']}"}},
            {"type":"text","text":f"Texto OCR:\n{texto_completo}\n\nRetorne o JSON."}
        ]}]
    else:
        msgs = [{"role":"user","content":
            f"DOCUMENTO ({conteudo['tipo'].upper()}):\n\n{texto_completo}\n\nRetorne o JSON conforme instruído."}]

    cli  = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
    resp = cli.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role":"system","content":prompt_sys}] + msgs,
        temperature=0.0, max_tokens=4000,
    )
    raw = resp.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"): raw = raw[4:]

    dados = json.loads(raw)
    if not isinstance(dados.get("etapas"), list):
        dados["etapas"] = []
    return dados


# ══════════════════════════════════════════════════════
#  GERAÇÃO APR — MODELO GAROA
# ══════════════════════════════════════════════════════

def gerar_apr_garoa(dados: dict, texto_doc: str) -> list:
    from openai import OpenAI

    etapas_txt = ""
    for e in dados.get("etapas", []):
        etapas_txt += (
            f"\nEtapa {e.get('numero','')}: {e.get('etapa','')}\n"
            f"  Ferramentas: {e.get('ferramentas','(nenhuma)')}\n"
            f"  Perigos doc: {e.get('perigos','(gere você)')}\n"
            f"  Prevenção doc: {e.get('como_prevenir','(gere você)')}\n"
        )

    texto_ref = _preparar_texto(texto_doc, MAX_CHARS_GERACAO)

    prompt = f"""Você é especialista em segurança do trabalho (NRs brasileiras).
Gere uma APR com Matriz de Risco detalhada.

DOCUMENTO DE REFERÊNCIA:
{texto_ref}

DADOS: Empresa:{dados.get('empresa','')} | Atividade:{dados.get('atividade','')}

ETAPAS:
{etapas_txt}

Para cada etapa gere:
- ferramentas: equipamentos específicos usados naquela etapa
- perigos: mínimo 3, formato "1.Risco A\\n2.Risco B\\n3.Risco C"
- como_prevenir: mínimo 3, formato "1.Medida A\\n2.Medida B\\n3.Medida C"
- freq: A (Muito Provável) | B (Provável) | C (Pouco Provável) | D (Remota)
- sev: I (Baixa) | II (Moderada) | III (Séria) | IV (Crítica)
- classe: "Risco Baixo" | "Risco Moderado" | "Risco Sério" | "Risco Crítico"

Retorne SOMENTE JSON:
{{"etapas":[{{"numero":1,"etapa":"...","ferramentas":"...","perigos":"1.X\\n2.Y\\n3.Z","como_prevenir":"1.A\\n2.B\\n3.C","freq":"B","sev":"II","classe":"Risco Moderado"}}]}}"""

    cli  = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
    resp = cli.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role":"user","content":prompt}],
        temperature=0.2, max_tokens=6000,
    )
    raw = resp.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"): raw = raw[4:]
    return json.loads(raw).get("etapas", [])


# ══════════════════════════════════════════════════════
#  GERAÇÃO APR — MODELO FR-EHS-04-01
# ══════════════════════════════════════════════════════

def gerar_apr_frehs(dados: dict, texto_doc: str) -> list:
    from openai import OpenAI

    etapas_txt = ""
    for e in dados.get("etapas", []):
        etapas_txt += (
            f"\nEtapa {e.get('numero','')}: {e.get('etapa','')}\n"
            f"  Ferramentas: {e.get('ferramentas','(nenhuma)')}\n"
            f"  Perigos doc: {e.get('perigos','(gere você)')}\n"
            f"  Prevenção doc: {e.get('como_prevenir','(gere você)')}\n"
        )

    texto_ref = _preparar_texto(texto_doc, MAX_CHARS_GERACAO)

    prompt = f"""Você é especialista em segurança do trabalho (NRs brasileiras).
Gere uma APR FR-EHS-04-01 detalhada.

DOCUMENTO DE REFERÊNCIA:
{texto_ref}

DADOS: Empresa:{dados.get('empresa','')} | Atividade:{dados.get('atividade','')}

ETAPAS:
{etapas_txt}

Para cada etapa gere:
- ferramentas: equipamentos específicos
- perigos: mínimo 3, formato "1.Risco A\\n2.Risco B\\n3.Risco C"
- como_prevenir: mínimo 3, formato "1.Medida A\\n2.Medida B\\n3.Medida C" (cite EPIs e NRs)

Retorne SOMENTE JSON:
{{"etapas":[{{"numero":1,"etapa":"...","ferramentas":"...","perigos":"1.X\\n2.Y\\n3.Z","como_prevenir":"1.A\\n2.B\\n3.C"}}]}}"""

    cli  = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com")
    resp = cli.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role":"user","content":prompt}],
        temperature=0.2, max_tokens=6000,
    )
    raw = resp.choices[0].message.content.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"): raw = raw[4:]
    return json.loads(raw).get("etapas", [])


# ══════════════════════════════════════════════════════
#  HELPERS EXCEL
# ══════════════════════════════════════════════════════

from openpyxl import Workbook as _Workbook
from openpyxl.styles import (Font as _Font, Alignment as _Alignment,
                              PatternFill as _PatternFill,
                              Border as _Border, Side as _Side)
from openpyxl.worksheet.page import PageMargins as _PageMargins

def _side(style="thin", color="595959"):
    return _Side(border_style=style, color=color)

def _border(style="thin", color="595959"):
    s = _side(style, color)
    return _Border(left=s, right=s, top=s, bottom=s)

def _fill(hex_color):
    return _PatternFill("solid", fgColor=hex_color)

def _font(bold=False, size=11, color="000000", name="Calibri"):
    return _Font(bold=bold, size=size, color=color, name=name)

def _aln(h="left", v="center", wrap=True):
    return _Alignment(horizontal=h, vertical=v, wrap_text=wrap)

def _apply(ws, row, col, value="", bold=False, size=11, color="000000",
           fill=None, halign="left", valign="center", wrap=True, border=True):
    cell = ws.cell(row=row, column=col)
    if value != "":
        cell.value = value
    cell.font      = _font(bold=bold, size=size, color=color)
    cell.alignment = _aln(halign, valign, wrap)
    if fill:
        cell.fill = _fill(fill)
    if border:
        cell.border = _border()
    return cell

def _merge_apply(ws, r1, c1, r2, c2, value="", bold=False, size=11,
                 color="000000", fill=None, halign="center", valign="center",
                 wrap=True, border=True):
    ws.merge_cells(start_row=r1, start_column=c1, end_row=r2, end_column=c2)
    cell = ws.cell(row=r1, column=c1)
    if value != "":
        cell.value = value
    cell.font      = _font(bold=bold, size=size, color=color)
    cell.alignment = _aln(halign, valign, wrap)
    if fill:
        cell.fill = _fill(fill)
    if border:
        bdr = _border()
        for r in range(r1, r2 + 1):
            for c in range(c1, c2 + 1):
                ws.cell(row=r, column=c).border = bdr
    return cell

def _risk_fill(classe: str):
    c = (classe or "").strip().lower()
    if "crítico" in c or "critico" in c:
        return "FF0000", "FFFFFF", True
    elif "sério" in c or "serio" in c:
        return "F4B084", "000000", False
    elif "moderado" in c:
        return "FFF2CC", "000000", False
    elif "baixo" in c:
        return "C6EFCE", "000000", False
    else:
        return "FFFFFF", "000000", False

def _clean_lines(raw):
    lines = [l.strip() for l in (raw or "").replace("•", "").split("\n") if l.strip()]
    return [_re.sub(r'^[\d]+[\.\)]\s*', '', l) for l in lines]


# ══════════════════════════════════════════════════════
#  EXPORTAÇÃO EXCEL — MODELO GAROA
# ══════════════════════════════════════════════════════

def exportar_excel_garoa(dados, etapas, caminho):
    wb = _Workbook()
    ws = wb.active
    ws.title = "APR"

    ws.page_setup.orientation = "landscape"
    ws.page_setup.paperSize   = 9
    ws.page_setup.fitToPage   = True
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 0
    ws.sheet_view.zoomScale   = 85
    ws.page_margins = _PageMargins(left=0.4, right=0.4, top=0.4, bottom=0.4,
                                   header=0.3, footer=0.3)

    ws.column_dimensions["A"].width = 6
    ws.column_dimensions["B"].width = 32.3
    ws.column_dimensions["C"].width = 44.0
    ws.column_dimensions["D"].width = 19.4
    ws.column_dimensions["E"].width = 14.6
    ws.column_dimensions["F"].width = 30.4
    ws.column_dimensions["G"].width = 53.4

    AZUL  = "1F4E78"
    CINZA = "D9D9D9"
    BRANC = "FFFFFF"
    row   = 1

    _merge_apply(ws, row, 1, row, 3, "FREQUÊNCIA",  bold=True, size=11, fill=CINZA)
    _merge_apply(ws, row, 4, row, 7, "SEVERIDADE",  bold=True, size=11, fill=CINZA)
    ws.row_dimensions[row].height = 18; row += 1

    freq_data = [
        ("A", "Muito provável",    "Evento com mais de uma ocorrência esperada ao longo da execução da atividade."),
        ("B", "Provável",          "Evento com pelo menos uma ocorrência esperada ao longo da execução da atividade."),
        ("C", "Pouco Provável",    "Evento com pelo menos uma ocorrência esperada ao longo de muitas execuções da atividade."),
        ("D", "Remota",            "Evento com muito baixa probabilidade de ocorrência ao longo de muitas execuções da atividade."),
    ]
    sev_data = [
        ("I",   "Baixa",    "- Danos insignificantes a funcionários ou público externo\n- Danos insignificantes ao meio ambiente\n- Danos insignificantes ao patrimônio"),
        ("II",  "Moderada", "Lesões leves em funcionários ou público externo\nDanos moderados ao meio ambiente\nDanos ao patrimônio de baixo valor"),
        ("III", "Séria",    "Lesões sérias em funcionários ou público externo\nDanos sérios ao meio ambiente\nDanos ao patrimônio de valor significativo"),
        ("IV",  "Critica",  "Mortes ou lesões graves em funcionários ou público externo\nDanos graves ao meio ambiente\nDanos de grande valor ao patrimônio"),
    ]
    row_h = [43.5, 30.0, 60.0, 45.0]
    for i, (letra, nome, desc) in enumerate(freq_data):
        ws.row_dimensions[row].height = row_h[i]
        _apply(ws, row, 1, letra, bold=True,  halign="center")
        _apply(ws, row, 2, nome,  bold=False, halign="center")
        _apply(ws, row, 3, desc,  bold=False, halign="center", wrap=True)
        sv = sev_data[i]
        _apply(ws, row, 4, sv[0], bold=True,  halign="center")
        _apply(ws, row, 5, sv[1], bold=True,  halign="center")
        _merge_apply(ws, row, 6, row, 7, sv[2], bold=False, halign="center", wrap=True)
        row += 1

    ws.row_dimensions[row].height = 9.75; row += 1

    _merge_apply(ws, row, 1, row+1, 2, "CLASSIFICAÇÃO DE RISCO", bold=False, halign="center")
    _merge_apply(ws, row, 3, row, 7, "Severidade", bold=False, halign="center", border=False)
    ws.row_dimensions[row].height = 15; row += 1

    ws.row_dimensions[row].height = 18.75
    for c, txt in enumerate(["I- Baixa","II - Moderada","III - Séria","IV - Crítica"], 3):
        _apply(ws, row, c, txt, bold=True, color=BRANC, fill=AZUL, halign="center")
    row += 1

    ws.merge_cells(start_row=row, start_column=1, end_row=row+3, end_column=1)
    ca = ws.cell(row=row, column=1, value="Freq.")
    ca.font = _font(bold=True); ca.alignment = _aln("center","center"); ca.border = _border()

    risk_rows = [
        ("A - Muito Provável", [("Risco Moderado","FFF2CC"),("Risco Sério","F4B084"),("Risco Crítico","FF0000"),("Risco Crítico","FF0000")]),
        ("B- Provável",        [("Risco Baixo","C6EFCE"),  ("Risco Moderado","FFF2CC"),("Risco Sério","FF0000"),  ("Risco Crítico","FF0000")]),
        ("C - Pouco Provável", [("Risco Baixo","C6EFCE"),  ("Risco Baixo","C6EFCE"),  ("Risco Moderado","FFF2CC"),("Risco Sério","FF0000")]),
        ("D - Remota",         [("Risco Baixo","C6EFCE"),  ("Risco Baixo","C6EFCE"),  ("Risco Baixo","C6EFCE"),  ("Risco Moderado","F4B084")]),
    ]
    tc_map = {"FF0000":("FFFFFF",True),"F4B084":("000000",False),"FFF2CC":("000000",False),"C6EFCE":("000000",False)}
    for r_data in risk_rows:
        ws.row_dimensions[row].height = 15
        _apply(ws, row, 2, r_data[0], bold=False)
        for ci, (txt, fhex) in enumerate(r_data[1], 3):
            tcol, tbold = tc_map.get(fhex, ("000000", False))
            _apply(ws, row, ci, txt, bold=tbold, color=tcol, fill=fhex, halign="center")
        row += 1

    ws.row_dimensions[row].height = 9.75; row += 1

    _merge_apply(ws, row, 1, row, 7, "PARTICIPANTE DA ANÁLISE", bold=True, halign="center")
    ws.row_dimensions[row].height = 15; row += 1

    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
    ws.merge_cells(start_row=row, start_column=4, end_row=row, end_column=5)
    ws.merge_cells(start_row=row, start_column=6, end_row=row, end_column=7)
    for col, val in [(1,"DATA"),(3,"NOME"),(4,"FUNÇÃO"),(6,"ASSINATURA")]:
        ws.cell(row=row, column=col, value=val).font = _font(bold=True)
        ws.cell(row=row, column=col).alignment = _aln("center","center")
        ws.cell(row=row, column=col).border    = _border()
    for col in [2, 5, 7]:
        ws.cell(row=row, column=col).border = _border()
    ws.row_dimensions[row].height = 15; row += 1

    for _ in range(3):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
        ws.merge_cells(start_row=row, start_column=4, end_row=row, end_column=5)
        ws.merge_cells(start_row=row, start_column=6, end_row=row, end_column=7)
        for col in range(1, 8):
            ws.cell(row=row, column=col).border    = _border()
            ws.cell(row=row, column=col).alignment = _aln()
        ws.row_dimensions[row].height = 18; row += 1

    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)
    ws.merge_cells(start_row=row, start_column=4, end_row=row, end_column=7)
    ws.cell(row=row, column=1, value=f"SERVIÇO A EXECUTAR: {dados.get('atividade','')}").font = _font(bold=True)
    ws.cell(row=row, column=1).alignment = _aln("left","center", wrap=True); ws.cell(row=row, column=1).border = _border()
    ws.cell(row=row, column=4, value=f"ENCARREGADO: {dados.get('responsavel','')}").font = _font(bold=True)
    ws.cell(row=row, column=4).alignment = _aln("left","center"); ws.cell(row=row, column=4).border = _border()
    for col in [2, 3, 5, 6, 7]: ws.cell(row=row, column=col).border = _border()
    ws.row_dimensions[row].height = 19; row += 1

    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)
    ws.merge_cells(start_row=row, start_column=4, end_row=row, end_column=5)
    ws.merge_cells(start_row=row, start_column=6, end_row=row, end_column=7)
    ws.cell(row=row, column=1, value=f"LOCAL DA OBRA: {dados.get('empresa','')}").font = _font()
    ws.cell(row=row, column=1).alignment = _aln("left","center"); ws.cell(row=row, column=1).border = _border()
    ws.cell(row=row, column=4, value=f"Nº APR: {dados.get('num_apr','')}").font = _font()
    ws.cell(row=row, column=4).alignment = _aln("left","center"); ws.cell(row=row, column=4).border = _border()
    ws.cell(row=row, column=6, value=f"DATA: {dados.get('data_inicio','')}").font = _font(bold=True)
    ws.cell(row=row, column=6).alignment = _aln("left","center"); ws.cell(row=row, column=6).border = _border()
    for col in [2, 3, 5, 7]: ws.cell(row=row, column=col).border = _border()
    ws.row_dimensions[row].height = 21.75; row += 1

    ws.row_dimensions[row].height = 23.25
    for c, txt in enumerate(["Nº","ETAPA DA ATIVIDADE","PERIGOS E RISCOS POTENCIAIS",
                              "FREQ.","SEV.","CLASSE DE RISCO",
                              "MEDIDAS DE CONTROLE / RECOMENDAÇÕES"], 1):
        _apply(ws, row, c, txt, bold=True, color=BRANC, fill=AZUL, halign="center")
    row += 1

    for etapa in etapas:
        num        = str(etapa.get("numero", ""))
        etapa_nome = etapa.get("etapa", "")
        perigos    = _clean_lines(etapa.get("perigos", ""))
        medidas    = _clean_lines(etapa.get("como_prevenir", ""))
        freq_val   = etapa.get("freq",   "B")
        sev_val    = etapa.get("sev",    "II")
        classe_val = etapa.get("classe", "Risco Moderado")

        max_lines = max(len(perigos), len(medidas), 1)
        while len(perigos) < max_lines: perigos.append("")
        while len(medidas) < max_lines: medidas.append("")

        start_row = row
        end_row   = row + max_lines - 1

        for i in range(max_lines):
            r = row + i
            ws.cell(row=r, column=1, value=num if i == 0 else "")
            ws.cell(row=r, column=2, value=etapa_nome if i == 0 else "")
            for col in range(1, 8):
                ws.cell(row=r, column=col).border = _border()
            ws.row_dimensions[r].height = 30

        if max_lines > 1:
            ws.merge_cells(start_row=start_row, start_column=1, end_row=end_row, end_column=1)
            ws.merge_cells(start_row=start_row, start_column=2, end_row=end_row, end_column=2)

        ws.cell(row=start_row, column=1).font      = _font(size=11)
        ws.cell(row=start_row, column=1).alignment = _aln("center","center")
        ws.cell(row=start_row, column=2).font      = _font(size=12)
        ws.cell(row=start_row, column=2).alignment = _aln("center","center", wrap=True)

        rf, fc, fb = _risk_fill(classe_val)
        for i in range(max_lines):
            r = row + i
            _apply(ws, r, 3, perigos[i], halign="left",   wrap=True)
            _apply(ws, r, 4, freq_val,   halign="center")
            _apply(ws, r, 5, sev_val,    halign="center")
            _apply(ws, r, 6, classe_val, bold=fb, color=fc, fill=rf, halign="center")
            _apply(ws, r, 7, medidas[i], halign="left",   wrap=True)

        row = end_row + 1

    for _ in range(3):
        for c in range(1, 8):
            _apply(ws, row, c, "")
        ws.row_dimensions[row].height = 30; row += 1

    wb.save(caminho)


# ══════════════════════════════════════════════════════
#  EXPORTAÇÃO EXCEL — MODELO FR-EHS-04-01
# ══════════════════════════════════════════════════════

def exportar_excel_frehs(dados, etapas, caminho):
    wb = _Workbook()
    ws = wb.active
    ws.title = "APR"

    ws.page_setup.orientation = "landscape"
    ws.page_setup.paperSize   = 9
    ws.page_setup.fitToPage   = True
    ws.page_setup.fitToWidth  = 1

    ws.column_dimensions["A"].width = 5
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 22
    ws.column_dimensions["D"].width = 38
    ws.column_dimensions["E"].width = 38

    AZUL  = "1F4E79"
    CINZA = "D9D9D9"
    row   = 1

    ws.merge_cells(f"A{row}:D{row}")
    ws.cell(row=row, column=1, value="ANÁLISE PRELIMINAR DE RISCO - APR")
    ws.cell(row=row, column=1).font      = _font(bold=True, size=14)
    ws.cell(row=row, column=1).alignment = _aln("center","center")
    ws.cell(row=row, column=1).border    = _border()
    ws.cell(row=row, column=5, value="FR-EHS-04-01")
    ws.cell(row=row, column=5).font      = _font(bold=True)
    ws.cell(row=row, column=5).alignment = _aln("center","center")
    ws.cell(row=row, column=5).border    = _border()
    for col in [2,3,4]: ws.cell(row=row, column=col).border = _border()
    ws.row_dimensions[row].height = 28; row += 1

    ws.merge_cells(f"A{row}:B{row}")
    ws.cell(row=row, column=1, value=f"EMPRESA: {dados.get('empresa','')}")
    ws.cell(row=row, column=1).font = _font(bold=True); ws.cell(row=row, column=1).border = _border(); ws.cell(row=row, column=1).alignment = _aln()
    ws.cell(row=row, column=2).border = _border()
    ws.merge_cells(f"C{row}:D{row}")
    ws.cell(row=row, column=3, value=f"RESPONSÁVEL: {dados.get('responsavel','')}")
    ws.cell(row=row, column=3).font = _font(); ws.cell(row=row, column=3).border = _border(); ws.cell(row=row, column=3).alignment = _aln()
    ws.cell(row=row, column=4).border = _border()
    ws.cell(row=row, column=5, value=f"Nº APR: {dados.get('num_apr','')}")
    ws.cell(row=row, column=5).font = _font(); ws.cell(row=row, column=5).border = _border(); ws.cell(row=row, column=5).alignment = _aln()
    ws.row_dimensions[row].height = 16; row += 1

    ws.merge_cells(f"A{row}:B{row}")
    ws.cell(row=row, column=1, value=f"DATA DE INÍCIO: {dados.get('data_inicio','')}")
    ws.cell(row=row, column=1).font = _font(); ws.cell(row=row, column=1).border = _border(); ws.cell(row=row, column=1).alignment = _aln()
    ws.cell(row=row, column=2).border = _border()
    ws.merge_cells(f"C{row}:D{row}")
    ws.cell(row=row, column=3, value=f"DATA PREVISTA TÉRMINO: {dados.get('data_termino','')}")
    ws.cell(row=row, column=3).font = _font(); ws.cell(row=row, column=3).border = _border(); ws.cell(row=row, column=3).alignment = _aln()
    ws.cell(row=row, column=4).border = _border()
    ws.cell(row=row, column=5, value="REVISÃO: ___/___/______ A ___/___/______")
    ws.cell(row=row, column=5).font = _font(); ws.cell(row=row, column=5).border = _border(); ws.cell(row=row, column=5).alignment = _aln()
    ws.row_dimensions[row].height = 16; row += 1

    ws.merge_cells(f"A{row}:E{row}")
    ws.cell(row=row, column=1, value=f"ATIVIDADE: {dados.get('atividade','')}")
    ws.cell(row=row, column=1).font = _font(bold=True); ws.cell(row=row, column=1).border = _border(); ws.cell(row=row, column=1).alignment = _aln()
    for col in [2,3,4,5]: ws.cell(row=row, column=col).border = _border()
    ws.row_dimensions[row].height = 16; row += 1

    headers = ["Nº","ETAPA DA ATIVIDADE","FERRAMENTAS\nEQUIPAMENTOS",
               "PERIGO / RISCO","COMO PREVENIR?\nUTILIZAR CONTROLES PADRÃO"]
    for col, txt in enumerate(headers, 1):
        ws.cell(row=row, column=col, value=txt)
        ws.cell(row=row, column=col).font      = _font(bold=True, color="FFFFFF")
        ws.cell(row=row, column=col).fill      = _fill(AZUL)
        ws.cell(row=row, column=col).border    = _border()
        ws.cell(row=row, column=col).alignment = _aln("center","center")
    ws.row_dimensions[row].height = 28; row += 1

    for e in etapas:
        ws.cell(row=row, column=1, value=e.get("numero",""))
        ws.cell(row=row, column=1).alignment = _aln("center","top")
        ws.cell(row=row, column=2, value=e.get("etapa",""))
        ws.cell(row=row, column=2).alignment = _aln()
        ws.cell(row=row, column=3, value=e.get("ferramentas","N.A"))
        ws.cell(row=row, column=3).alignment = _aln()
        ws.cell(row=row, column=4, value=(e.get("perigos","") or "").replace("<br/>","\n"))
        ws.cell(row=row, column=4).alignment = _aln()
        ws.cell(row=row, column=5, value=(e.get("como_prevenir","") or "").replace("<br/>","\n"))
        ws.cell(row=row, column=5).alignment = _aln()
        for col in range(1, 6):
            ws.cell(row=row, column=col).border = _border()
            ws.cell(row=row, column=col).font   = _font(size=9)
        ws.row_dimensions[row].height = 70; row += 1

    wb.save(caminho)


# ══════════════════════════════════════════════════════
#  EXPORTAÇÃO PDF — MODELO GAROA
# ══════════════════════════════════════════════════════

def exportar_pdf_garoa(dados, etapas, caminho):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    PAGE = landscape(A4)
    doc  = SimpleDocTemplate(caminho, pagesize=PAGE,
                             rightMargin=1*cm, leftMargin=1*cm,
                             topMargin=1*cm, bottomMargin=1*cm)
    W = PAGE[0] - 2*cm

    AZUL   = colors.HexColor("#1F4E78")
    CINZA  = colors.HexColor("#D9D9D9")
    ESCURO = colors.HexColor("#595959")
    BRANC  = colors.white
    thin   = 0.5

    COR_RISCO = {
        "risco crítico":  (colors.HexColor("#FF0000"), BRANC,  True),
        "risco sério":    (colors.HexColor("#F4B084"), colors.black, False),
        "risco moderado": (colors.HexColor("#FFF2CC"), colors.black, False),
        "risco baixo":    (colors.HexColor("#C6EFCE"), colors.black, False),
    }

    def ps(name, size=8, bold=False, align=TA_LEFT, color=colors.black, leading=10):
        return ParagraphStyle(name,
            fontName="Helvetica-Bold" if bold else "Helvetica",
            fontSize=size, alignment=align, textColor=color,
            leading=leading, spaceAfter=0, spaceBefore=0)

    h_s  = ps("h",  8, bold=True,  align=TA_CENTER, color=BRANC)
    c_s  = ps("c",  7, bold=False, align=TA_LEFT,   color=colors.black)
    cc_s = ps("cc", 7, bold=False, align=TA_CENTER, color=colors.black)
    v_s  = ps("v",  8, bold=False, align=TA_LEFT,   color=colors.black)

    elems = []

    freq_data = [
        ("A","Muito Provável","Evento com mais de uma ocorrência esperada ao longo da execução da atividade."),
        ("B","Provável","Evento com pelo menos uma ocorrência esperada ao longo da execução da atividade."),
        ("C","Pouco Provável","Evento com pelo menos uma ocorrência esperada ao longo de muitas execuções da atividade."),
        ("D","Remota","Evento com muito baixa probabilidade de ocorrência ao longo de muitas execuções da atividade."),
    ]
    sev_data = [
        ("I","Baixa","Danos insignificantes a funcionários ou público externo\nDanos insignificantes ao meio ambiente\nDanos insignificantes ao patrimônio"),
        ("II","Moderada","Lesões leves em funcionários ou público externo\nDanos moderados ao meio ambiente\nDanos ao patrimônio de baixo valor"),
        ("III","Séria","Lesões sérias em funcionários ou público externo\nDanos sérios ao meio ambiente\nDanos ao patrimônio de valor significativo"),
        ("IV","Crítica","Mortes ou lesões graves em funcionários ou público externo\nDanos graves ao meio ambiente\nDanos de grande valor ao patrimônio"),
    ]

    leg_header = Table(
        [[Paragraph("FREQUÊNCIA", ps("fh",8,bold=True,align=TA_CENTER)),
          Paragraph("SEVERIDADE", ps("sh",8,bold=True,align=TA_CENTER))]],
        colWidths=[W*0.40, W*0.60])
    leg_header.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),thin,ESCURO),("LINEAFTER",(0,0),(0,-1),thin,ESCURO),
        ("BACKGROUND",(0,0),(-1,-1),CINZA),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
    ]))
    elems.append(leg_header)

    leg_rows = []
    for i in range(4):
        fl, fn, fd = freq_data[i]
        sl, sn, sd = sev_data[i]
        leg_rows.append([
            Paragraph(fl, ps(f"fl{i}",8,bold=True,align=TA_CENTER)),
            Paragraph(fn, cc_s),
            Paragraph(fd.replace("\n","<br/>"), c_s),
            Paragraph(sl, ps(f"sl{i}",8,bold=True,align=TA_CENTER)),
            Paragraph(sn, ps(f"sn{i}",8,bold=True,align=TA_CENTER)),
            Paragraph(sd.replace("\n","<br/>"), c_s),
        ])

    leg_table = Table(leg_rows, colWidths=[W*0.04, W*0.12, W*0.24, W*0.05, W*0.10, W*0.45])
    leg_table.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),thin,ESCURO),("INNERGRID",(0,0),(-1,-1),thin,ESCURO),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),("LEFTPADDING",(0,0),(-1,-1),3),
    ]))
    elems.append(leg_table)
    elems.append(Spacer(1, 0.10*cm))

    mat_col_w = [W*0.06, W*0.16, W*0.195, W*0.195, W*0.195, W*0.195]
    RB  = colors.HexColor("#C6EFCE")
    RM  = colors.HexColor("#FFF2CC")
    RS  = colors.HexColor("#F4B084")
    RC  = colors.HexColor("#FF0000")

    mat_rows = [
        [Paragraph("CLASSIFICAÇÃO\nDE RISCO", ps("cr",7,bold=True,align=TA_CENTER)),
         Paragraph("", c_s),
         Paragraph("Severidade", ps("sv",8,bold=True,align=TA_CENTER)),
         Paragraph("",c_s), Paragraph("",c_s), Paragraph("",c_s)],
        [Paragraph("",c_s), Paragraph("",c_s),
         Paragraph("I - Baixa",    ps("i1",7,bold=True,align=TA_CENTER,color=BRANC)),
         Paragraph("II - Moderada",ps("i2",7,bold=True,align=TA_CENTER,color=BRANC)),
         Paragraph("III - Séria",  ps("i3",7,bold=True,align=TA_CENTER,color=BRANC)),
         Paragraph("IV - Crítica", ps("i4",7,bold=True,align=TA_CENTER,color=BRANC))],
        [Paragraph("Freq.", ps("fq",7,bold=True,align=TA_CENTER)),
         Paragraph("A - Muito Provável", ps("fa",7,align=TA_LEFT)),
         Paragraph("Risco Moderado", ps("rm0",7,align=TA_CENTER)),
         Paragraph("Risco Sério",    ps("rs0",7,align=TA_CENTER)),
         Paragraph("Risco Crítico",  ps("rc0",7,bold=True,align=TA_CENTER,color=BRANC)),
         Paragraph("Risco Crítico",  ps("rc1",7,bold=True,align=TA_CENTER,color=BRANC))],
        [Paragraph("",c_s),
         Paragraph("B - Provável", ps("fb",7,align=TA_LEFT)),
         Paragraph("Risco Baixo",    ps("rb1",7,align=TA_CENTER)),
         Paragraph("Risco Moderado", ps("rm1",7,align=TA_CENTER)),
         Paragraph("Risco Sério",    ps("rs1",7,bold=True,align=TA_CENTER,color=BRANC)),
         Paragraph("Risco Crítico",  ps("rc2",7,bold=True,align=TA_CENTER,color=BRANC))],
        [Paragraph("",c_s),
         Paragraph("C - Pouco Provável", ps("fc",7,align=TA_LEFT)),
         Paragraph("Risco Baixo",    ps("rb2",7,align=TA_CENTER)),
         Paragraph("Risco Baixo",    ps("rb3",7,align=TA_CENTER)),
         Paragraph("Risco Moderado", ps("rm2",7,align=TA_CENTER)),
         Paragraph("Risco Sério",    ps("rs2",7,bold=True,align=TA_CENTER,color=BRANC))],
        [Paragraph("",c_s),
         Paragraph("D - Remota", ps("fd",7,align=TA_LEFT)),
         Paragraph("Risco Baixo",    ps("rb4",7,align=TA_CENTER)),
         Paragraph("Risco Baixo",    ps("rb5",7,align=TA_CENTER)),
         Paragraph("Risco Baixo",    ps("rb6",7,align=TA_CENTER)),
         Paragraph("Risco Moderado", ps("rm3",7,align=TA_CENTER))],
    ]

    mat = Table(mat_rows, colWidths=mat_col_w)
    mat.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),thin,ESCURO),("INNERGRID",(0,0),(-1,-1),thin,ESCURO),
        ("SPAN",(0,0),(0,1)),("SPAN",(2,0),(5,0)),("SPAN",(0,2),(0,5)),
        ("BACKGROUND",(2,1),(5,1), AZUL),
        ("BACKGROUND",(2,2),(2,2), RM),("BACKGROUND",(3,2),(3,2), RS),
        ("BACKGROUND",(4,2),(5,2), RC),("BACKGROUND",(2,3),(2,3), RB),
        ("BACKGROUND",(3,3),(3,3), RM),("BACKGROUND",(4,3),(5,3), RC),
        ("BACKGROUND",(2,4),(3,4), RB),("BACKGROUND",(4,4),(4,4), RM),
        ("BACKGROUND",(5,4),(5,4), RC),("BACKGROUND",(2,5),(4,5), RB),
        ("BACKGROUND",(5,5),(5,5), RS),
        ("TEXTCOLOR",(4,2),(5,2), BRANC),("TEXTCOLOR",(4,3),(5,3), BRANC),
        ("TEXTCOLOR",(5,4),(5,4), BRANC),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",(0,0),(-1,-1),3),
    ]))
    elems.append(mat)
    elems.append(Spacer(1, 0.10*cm))

    part_title = Table(
        [[Paragraph("PARTICIPANTE DA ANÁLISE", ps("pt",8,bold=True,align=TA_CENTER))]],
        colWidths=[W])
    part_title.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),thin,ESCURO),("BACKGROUND",(0,0),(-1,-1),CINZA),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
    ]))
    elems.append(part_title)

    part_cw = [W*0.15, W*0.35, W*0.25, W*0.25]
    part_header = Table(
        [[Paragraph("DATA",h_s), Paragraph("NOME",h_s),
          Paragraph("FUNÇÃO",h_s), Paragraph("ASSINATURA",h_s)]],
        colWidths=part_cw)
    part_header.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),thin,ESCURO),("INNERGRID",(0,0),(-1,-1),thin,ESCURO),
        ("BACKGROUND",(0,0),(-1,-1),AZUL),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
    ]))
    elems.append(part_header)

    part_body = Table([[Paragraph("",c_s)]*4 for _ in range(3)], colWidths=part_cw)
    part_body.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),thin,ESCURO),("INNERGRID",(0,0),(-1,-1),thin,ESCURO),
        ("TOPPADDING",(0,0),(-1,-1),10),("BOTTOMPADDING",(0,0),(-1,-1),10),
    ]))
    elems.append(part_body)
    elems.append(Spacer(1, 0.10*cm))

    serv_data = Table([
        [Paragraph(f"<b>SERVIÇO A EXECUTAR:</b> {dados.get('atividade','')}", v_s),
         Paragraph(f"<b>ENCARREGADO:</b> {dados.get('responsavel','')}", v_s)],
        [Paragraph(f"<b>LOCAL DA OBRA:</b> {dados.get('empresa','')} "
                   f"    <b>Nº APR:</b> {dados.get('num_apr','')} "
                   f"    <b>DATA:</b> {dados.get('data_inicio','')}", v_s),
         Paragraph("", v_s)],
    ], colWidths=[W*0.55, W*0.45])
    serv_data.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),thin,ESCURO),("INNERGRID",(0,0),(-1,-1),thin,ESCURO),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",(0,0),(-1,-1),4),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
    ]))
    elems.append(serv_data)
    elems.append(Spacer(1, 0.15*cm))

    col_w = [W*0.04, W*0.20, W*0.28, W*0.05, W*0.05, W*0.12, W*0.26]
    apr_rows = [[
        Paragraph("Nº", h_s),
        Paragraph("ETAPA DA ATIVIDADE", h_s),
        Paragraph("PERIGOS E RISCOS POTENCIAIS", h_s),
        Paragraph("FREQ.", h_s),
        Paragraph("SEV.", h_s),
        Paragraph("CLASSE DE RISCO", h_s),
        Paragraph("MEDIDAS DE CONTROLE / RECOMENDAÇÕES", h_s),
    ]]

    cmds = [
        ("BOX",(0,0),(-1,-1),1,ESCURO),("INNERGRID",(0,0),(-1,-1),thin,ESCURO),
        ("BACKGROUND",(0,0),(-1,0),AZUL),("VALIGN",(0,0),(-1,-1),"TOP"),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("LEFTPADDING",(0,0),(-1,-1),3),
        ("ALIGN",(0,0),(0,-1),"CENTER"),("ALIGN",(3,0),(5,-1),"CENTER"),
    ]

    for etapa in etapas:
        perigos  = _clean_lines(etapa.get("perigos",""))
        medidas  = _clean_lines(etapa.get("como_prevenir",""))
        freq_val = etapa.get("freq","B")
        sev_val  = etapa.get("sev","II")
        classe   = etapa.get("classe","Risco Moderado")
        cor_bg, cor_txt, bold_txt = COR_RISCO.get(classe.lower().strip(), (colors.white, colors.black, False))

        max_l = max(len(perigos), len(medidas), 1)
        while len(perigos) < max_l: perigos.append("")
        while len(medidas) < max_l: medidas.append("")

        start = len(apr_rows)
        for i in range(max_l):
            apr_rows.append([
                Paragraph(str(etapa.get("numero","")) if i==0 else "", ps(f"n{start}{i}",8,align=TA_CENTER)),
                Paragraph(etapa.get("etapa","") if i==0 else "", c_s),
                Paragraph(perigos[i], c_s),
                Paragraph(freq_val, ps(f"f{start}{i}",7,align=TA_CENTER)),
                Paragraph(sev_val,  ps(f"s{start}{i}",7,align=TA_CENTER)),
                Paragraph(classe,   ps(f"cl{start}{i}",7,bold=bold_txt,align=TA_CENTER,color=cor_txt)),
                Paragraph(medidas[i], c_s),
            ])
            ri = len(apr_rows) - 1
            cmds.append(("BACKGROUND",(2,ri),(6,ri), cor_bg))

        end = len(apr_rows) - 1
        if max_l > 1:
            cmds.append(("SPAN",(0,start),(0,end)))
            cmds.append(("SPAN",(1,start),(1,end)))

    for _ in range(3):
        apr_rows.append([Paragraph("",c_s)]*7)

    apr_table = Table(apr_rows, colWidths=col_w, repeatRows=1)
    apr_table.setStyle(TableStyle(cmds))
    elems.append(apr_table)
    elems.append(Spacer(1, 0.3*cm))

    # ── Orientações de emergência ──
    t_emerg = Table([[Paragraph(
        "ORIENTAÇÕES EM CASO DE EMERGÊNCIA OU INCIDENTES: PROCURAR EQUIPE DE TÉCNICOS DE "
        "SEGURANÇA PARA ENCAMINHAMENTO. NA IMPOSSIBILIDADE, PEDIR AJUDA PARA A PESSOA MAIS PRÓXIMA.",
        ps("warn", 7, bold=True))]], colWidths=[W])
    t_emerg.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),thin,ESCURO),
        ("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#FFF2CC")),
        ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),("LEFTPADDING",(0,0),(-1,-1),4),
    ]))
    elems.append(t_emerg)
    elems.append(Spacer(1, 0.2*cm))

    # ── Encerramento da APR ──
    enc_title = Table(
        [[Paragraph("ENCERRAMENTO DA APR", ps("et", 8, bold=True, align=TA_CENTER))]],
        colWidths=[W])
    enc_title.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),thin,ESCURO),
        ("BACKGROUND",(0,0),(-1,-1),CINZA),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
    ]))
    elems.append(enc_title)

    enc_cw = [W*0.25, W*0.35, W*0.20, W*0.20]
    enc_header = Table(
        [[Paragraph("RESPONSÁVEL", h_s), Paragraph("FUNÇÃO", h_s),
          Paragraph("DATA", h_s), Paragraph("ASSINATURA", h_s)]],
        colWidths=enc_cw)
    enc_header.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),thin,ESCURO),("INNERGRID",(0,0),(-1,-1),thin,ESCURO),
        ("BACKGROUND",(0,0),(-1,-1),AZUL),
        ("ALIGN",(0,0),(-1,-1),"CENTER"),("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),
    ]))
    elems.append(enc_header)

    enc_body = Table([[Paragraph("",c_s)]*4 for _ in range(2)], colWidths=enc_cw)
    enc_body.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),thin,ESCURO),("INNERGRID",(0,0),(-1,-1),thin,ESCURO),
        ("TOPPADDING",(0,0),(-1,-1),12),("BOTTOMPADDING",(0,0),(-1,-1),12),
    ]))
    elems.append(enc_body)

    doc.build(elems)


# ══════════════════════════════════════════════════════
#  EXPORTAÇÃO PDF — MODELO FR-EHS-04-01
# ══════════════════════════════════════════════════════

def exportar_pdf_frehs(dados, etapas, caminho):
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    PAGE = landscape(A4)
    doc  = SimpleDocTemplate(caminho, pagesize=PAGE,
                             rightMargin=1*cm, leftMargin=1*cm,
                             topMargin=1*cm, bottomMargin=1*cm)
    W = PAGE[0] - 2*cm

    def ps(name, size=8, bold=False, align=TA_LEFT, color=colors.black, leading=11):
        return ParagraphStyle(name, fontName="Helvetica-Bold" if bold else "Helvetica",
                              fontSize=size, alignment=align, textColor=color,
                              leading=leading, spaceAfter=0, spaceBefore=0)

    CINZA  = colors.HexColor("#D9D9D9")
    ESCURO = colors.HexColor("#404040")
    AZUL   = colors.HexColor("#1F4E79")
    BORDA  = colors.HexColor("#888888")
    thin   = 0.5
    elems  = []

    valor_s  = ps("valor",  8)
    header_s = ps("header", 8, bold=True, align=TA_CENTER, color=colors.white)
    cell_s   = ps("cell",   7, leading=10)
    warn_s   = ps("warn",   7, bold=True)
    small_s  = ps("small",  6, align=TA_CENTER)

    t = Table([[Paragraph("ANÁLISE PRELIMINAR DE RISCO - APR", ps("titulo",16,bold=True,align=TA_CENTER)),
                [Paragraph("FR-EHS-04-01", ps("cod",8,bold=True,align=TA_CENTER)),
                 Paragraph("PÁGINA:", small_s)]]],
              colWidths=[W*0.82, W*0.18])
    t.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),1,ESCURO),("LINEAFTER",(0,0),(0,-1),1,ESCURO),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6),
    ]))
    elems.append(t)

    t = Table([[Paragraph(f"<b>EMPRESA:</b> {dados.get('empresa','')}", valor_s),
                Paragraph(f"<b>RESPONSÁVEL:</b> {dados.get('responsavel','')}", valor_s),
                Paragraph(f"<b>Nº APR:</b> {dados.get('num_apr','')}", valor_s)]],
              colWidths=[W*0.35, W*0.45, W*0.20])
    t.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),thin,ESCURO),("LINEBEFORE",(1,0),(2,-1),thin,ESCURO),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),("LEFTPADDING",(0,0),(-1,-1),4),
    ]))
    elems.append(t)

    t = Table([[Paragraph(f"<b>DATA DE INÍCIO:</b> {dados.get('data_inicio','')}", valor_s),
                Paragraph(f"<b>DATA PREVISTA TÉRMINO:</b> {dados.get('data_termino','')}", valor_s),
                Paragraph("<b>REVISÃO:</b> ___/___/______", valor_s)]],
              colWidths=[W*0.25, W*0.35, W*0.40])
    t.setStyle(TableStyle([
        ("BOX",(0,0),(-1,-1),thin,ESCURO),("LINEBEFORE",(1,0),(2,-1),thin,ESCURO),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),("LEFTPADDING",(0,0),(-1,-1),4),
    ]))
    elems.append(t)

    t = Table([[Paragraph(f"<b>ATIVIDADE:</b> {dados.get('atividade','')}", valor_s)]], colWidths=[W])
    t.setStyle(TableStyle([("BOX",(0,0),(-1,-1),thin,ESCURO),
        ("TOPPADDING",(0,0),(-1,-1),3),("BOTTOMPADDING",(0,0),(-1,-1),3),("LEFTPADDING",(0,0),(-1,-1),4)]))
    elems.append(t)
    elems.append(Spacer(1, 0.2*cm))

    col_w = [W*0.04, W*0.20, W*0.14, W*0.31, W*0.31]
    rows  = [[Paragraph("Nº",header_s), Paragraph("ETAPA DA ATIVIDADE",header_s),
              Paragraph("FERRAMENTAS\nEQUIPAMENTOS",header_s),
              Paragraph("PERIGO / RISCO",header_s),
              Paragraph("COMO PREVENIR?\nUTILIZAR CONTROLES PADRÃO",header_s)]]

    for e in etapas:
        rows.append([
            Paragraph(str(e.get("numero","")), ps("n",8,align=TA_CENTER)),
            Paragraph((e.get("etapa","") or "").replace("\n","<br/>"), cell_s),
            Paragraph((e.get("ferramentas","") or "N.A").replace("\n","<br/>"), cell_s),
            Paragraph((e.get("perigos","") or "").replace("\n","<br/>"), cell_s),
            Paragraph((e.get("como_prevenir","") or "").replace("\n","<br/>"), cell_s),
        ])

    cmds = [
        ("BOX",(0,0),(-1,-1),1,ESCURO),("INNERGRID",(0,0),(-1,-1),thin,BORDA),
        ("BACKGROUND",(0,0),(-1,0),AZUL),("VALIGN",(0,0),(-1,-1),"TOP"),
        ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),
        ("LEFTPADDING",(0,0),(-1,-1),3),("ALIGN",(0,0),(0,-1),"CENTER"),
    ]
    t = Table(rows, colWidths=col_w, repeatRows=1)
    t.setStyle(TableStyle(cmds))
    elems.append(t)
    elems.append(Spacer(1, 0.3*cm))

    t = Table([[Paragraph(
        "ORIENTAÇÕES EM CASO DE EMERGÊNCIA OU INCIDENTES: PROCURAR EQUIPE DE TÉCNICOS DE "
        "SEGURANÇA PARA ENCAMINHAMENTO. NA IMPOSSIBILIDADE, PEDIR AJUDA PARA A PESSOA MAIS PRÓXIMA.",
        warn_s)]], colWidths=[W])
    t.setStyle(TableStyle([("BOX",(0,0),(-1,-1),thin,ESCURO),
        ("BACKGROUND",(0,0),(-1,-1),colors.HexColor("#FFF2CC")),
        ("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4),("LEFTPADDING",(0,0),(-1,-1),4)]))
    elems.append(t)
    doc.build(elems)


# ══════════════════════════════════════════════════════
#  PIPELINE PRINCIPAL
# ══════════════════════════════════════════════════════

def pipeline(job_id: str, arquivo_path: str, modelo: str):
    job = jobs[job_id]

    def evento(tipo: str, dados: dict):
        job["eventos"].append({"tipo": tipo, "dados": dados})

    try:
        evento("step", {"step": 0, "status": "active"})
        conteudo = carregar_arquivo(arquivo_path)
        evento("step", {"step": 0, "status": "done"})

        evento("step", {"step": 1, "status": "active"})
        dados = extrair_info(conteudo, modelo)
        evento("step", {"step": 1, "status": "done"})

        evento("dados", {
            "empresa":      dados.get("empresa",""),
            "responsavel":  dados.get("responsavel",""),
            "num_apr":      dados.get("num_apr",""),
            "atividade":    dados.get("atividade",""),
            "data_inicio":  dados.get("data_inicio",""),
            "data_termino": dados.get("data_termino",""),
        })

        evento("step", {"step": 2, "status": "active"})

        if modelo == "garoa":
            etapas = gerar_apr_garoa(dados, conteudo["texto"])
        else:
            etapas = gerar_apr_frehs(dados, conteudo["texto"])

        evento("step", {"step": 2, "status": "done"})
        evento("step", {"step": 3, "status": "active"})
        evento("step", {"step": 3, "status": "done"})
        evento("step", {"step": 4, "status": "active"})

        pdf_path = TEMP_DIR / f"{job_id}.pdf"
        xls_path = TEMP_DIR / f"{job_id}.xlsx"

        if modelo == "garoa":
            exportar_pdf_garoa(dados, etapas, str(pdf_path))
            exportar_excel_garoa(dados, etapas, str(xls_path))
        else:
            exportar_pdf_frehs(dados, etapas, str(pdf_path))
            exportar_excel_frehs(dados, etapas, str(xls_path))

        evento("step", {"step": 4, "status": "done"})
        evento("pronto", {
            "pdf_url":  f"/download/{job_id}/pdf",
            "xlsx_url": f"/download/{job_id}/xlsx",
            "etapas":   len(etapas),
        })
        job["status"] = "done"

    except Exception as exc:
        import traceback
        evento("erro", {"mensagem": str(exc), "detalhe": traceback.format_exc()})
        job["status"] = "error"
    finally:
        try:
            Path(arquivo_path).unlink(missing_ok=True)
        except Exception:
            pass


# ══════════════════════════════════════════════════════
#  ENDPOINTS
# ══════════════════════════════════════════════════════

@app.get("/")
async def health():
    return {"status": "ok", "service": "APR Generator API"}


@app.post("/processar")
async def processar(
    arquivo: UploadFile = File(...),
    modelo:  str        = Form("garoa"),
):
    import threading

    ext      = Path(arquivo.filename).suffix.lower()
    job_id   = str(uuid.uuid4())
    tmp_path = TEMP_DIR / f"{job_id}_upload{ext}"

    contents = await arquivo.read()
    tmp_path.write_bytes(contents)

    jobs[job_id] = {
        "status": "running",
        "eventos": [],
        "cursor": 0,
    }

    t = threading.Thread(
        target=pipeline,
        args=(job_id, str(tmp_path), modelo),
        daemon=True,
    )
    t.start()

    return {"job_id": job_id}


@app.get("/stream/{job_id}")
async def stream(job_id: str):
    if job_id not in jobs:
        raise HTTPException(status_code=404, detail="Job não encontrado")

    async def event_generator():
        job    = jobs[job_id]
        cursor = 0
        while True:
            eventos = job["eventos"]
            while cursor < len(eventos):
                ev = eventos[cursor]
                cursor += 1
                yield f"event: {ev['tipo']}\ndata: {json.dumps(ev['dados'], ensure_ascii=False)}\n\n"

            if job["status"] in ("done", "error") and cursor >= len(job["eventos"]):
                break

            await asyncio.sleep(0.2)

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control":               "no-cache",
            "X-Accel-Buffering":           "no",
            "Access-Control-Allow-Origin": "*",
        },
    )


# ══════════════════════════════════════════════════════
#  DOWNLOAD — GET + HEAD com headers corretos
# ══════════════════════════════════════════════════════

@app.api_route("/download/{job_id}/{tipo}", methods=["GET", "HEAD"])
async def download(job_id: str, tipo: str, request: Request):
    if tipo == "pdf":
        path  = TEMP_DIR / f"{job_id}.pdf"
        media = "application/pdf"
        nome  = "APR.pdf"
    elif tipo == "xlsx":
        path  = TEMP_DIR / f"{job_id}.xlsx"
        media = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        nome  = "APR.xlsx"
    else:
        raise HTTPException(status_code=400, detail="Tipo inválido. Use 'pdf' ou 'xlsx'.")

    if not path.exists():
        raise HTTPException(status_code=404, detail="Arquivo não encontrado ou ainda sendo gerado.")

    headers = {
        "Content-Disposition": f'attachment; filename="{nome}"',
        "Access-Control-Allow-Origin":  "*",
        "Access-Control-Allow-Methods": "GET, HEAD, OPTIONS",
        "Access-Control-Allow-Headers": "*",
        "Cache-Control":                "no-cache, no-store, must-revalidate",
        "X-Frame-Options":              "ALLOWALL",
        "Content-Security-Policy":      "frame-ancestors *",
    }

    if request.method == "HEAD":
        return Response(status_code=200, media_type=media, headers=headers)

    content = path.read_bytes()
    return Response(content=content, media_type=media, headers=headers)


@app.options("/download/{job_id}/{tipo}")
async def download_options(job_id: str, tipo: str):
    return Response(
        status_code=200,
        headers={
            "Access-Control-Allow-Origin":  "*",
            "Access-Control-Allow-Methods": "GET, HEAD, OPTIONS",
            "Access-Control-Allow-Headers": "*",
        },
    )


# ══════════════════════════════════════════════════════
#  ENTRYPOINT
# ══════════════════════════════════════════════════════

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=False)